import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding/margins of a cell in twentieths of a point (dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Color palette
    PRIMARY_COLOR = RGBColor(31, 78, 121)     # Steel Blue
    SECONDARY_COLOR = RGBColor(112, 128, 144) # Slate Grey
    TEXT_COLOR = RGBColor(51, 51, 51)         # Dark Grey
    ACCENT_WARN = RGBColor(192, 57, 43)       # Rust Red
    
    # Base Style modifications
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = TEXT_COLOR
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Document Title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.space_before = Pt(6)
    run_title = title.add_run("Day 1 Data Ingestion & Quality Summary")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR
    
    # Subtitle/Metadata
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(24)
    run_meta = meta.add_run("Mutual Fund Analytics Capstone Project | Day 1 Work Review")
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = SECONDARY_COLOR
    
    # Intro
    p = doc.add_paragraph("Hey,")
    p = doc.add_paragraph(
        "Here is the complete summary of the work I did today for the Mutual Fund Analytics setup. "
        "I've got the project folder structure organized, the local environment set up, all the raw "
        "datasets ingested, and the live API connection working."
    )
    p = doc.add_paragraph(
        "I did a deep dive into the 10 CSV files to check for quality issues. There are a few dirty "
        "spots (duplicates, missing fields, and date format mismatches) that we will need to clean up "
        "first thing tomorrow. I’ve detailed everything below."
    )
    
    # Section 1
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    run = h1.add_run("1. Project Directory & Environment Layout")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "I set up the project folder structure inside the directory: "
    )
    p.add_run("C:\\Users\\jibum\\OneDrive\\Desktop\\Bluestock Internship").bold = True
    
    p = doc.add_paragraph("Here is the current layout and where we stand:")
    
    # Bullet points with manual indent
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("data/raw/").bold = True
    bp1.add_run(" - Active. This contains the 10 original local CSV datasets, plus the raw historical CSVs I downloaded from the API.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("data/processed/").bold = True
    bp2.add_run(" - Empty. This is intentional! It is currently an empty staging folder. I will be outputting our deduplicated, clean, and merged datasets here during tomorrow's cleaning phase.")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("notebooks/").bold = True
    bp3.add_run(" - Placeholder. Set up and ready for Jupyter/Google Colab notebooks for explorative analysis.")
    
    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.add_run("sql/").bold = True
    bp4.add_run(" - Placeholder. Ready to hold database table definitions, schemas, and staging queries.")
    
    bp5 = doc.add_paragraph(style='List Bullet')
    bp5.add_run("dashboard/").bold = True
    bp5.add_run(" - Empty. Another intentional placeholder. I'll be using this directory later to store the visual assets, styling sheets, and configuration files for the UI and dashboard widgets.")
    
    bp6 = doc.add_paragraph(style='List Bullet')
    bp6.add_run("reports/").bold = True
    bp6.add_run(" - Active. Where I'm saving our documentation, summaries, and data quality reports.")
    
    # Quick Note
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(8)
    p_note.add_run("Quick Note on Environment: ").bold = True
    p_note.add_run(
        "I successfully initialized the local Git repository and set up the remote pointing to our GitHub repository. "
        "Everything has been committed cleanly under the message \"Day 1: Data ingestion complete\"."
    )
    p_note = doc.add_paragraph(
        "I also created requirements.txt with all the specific versions we need for the analytical stack (Pandas, NumPy, "
        "Matplotlib, Seaborn, Plotly, SQLAlchemy, Requests, SciPy, and Jupyter)."
    )
    
    # Section 2
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    run = h2.add_run("2. Ingested Datasets & Structural Profiles")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "I wrote an ETL script (data_ingestion.py) to systematically load and inspect all 10 source CSV files using Pandas. "
        "Here is the exact profile of the data as it stands:"
    )
    
    # Table Creation
    headers = ["Dataset File Name", "Shape", "Target Primary Key", "Main Column Dtypes", "Quality Status"]
    data = [
        ["fund_master.csv", "12 x 6", "amfi_code", "int64, object (string)", "Dirty (Nulls & Duplicates)"],
        ["nav_history.csv", "101 x 3", "amfi_code + date", "int64, object (string), float64", "Dirty (Mixed date formats)"],
        ["investors.csv", "5 x 5", "investor_id", "object, object (string)", "Dirty (Has nulls)"],
        ["transactions.csv", "6 x 7", "transaction_id", "object, int64, float64", "Dirty (Nulls & Duplicates)"],
        ["fund_managers.csv", "4 x 4", "manager_id", "object, int64, object", "Clean"],
        ["portfolio_holdings.csv", "5 x 5", "holding_id", "object, int64, float64", "Clean"],
        ["amc_details.csv", "4 x 5", "amc_id", "object, int64, int64", "Clean"],
        ["expense_ratios.csv", "4 x 4", "amfi_code", "int64, float64, float64", "Clean"],
        ["benchmarks.csv", "4 x 3", "benchmark_id", "object, object, object", "Clean"],
        ["benchmark_history.csv", "20 x 3", "benchmark_id + date", "object, object, float64", "Clean"]
    ]
    
    table = doc.add_table(rows=1 + len(data), cols=5)
    table.style = 'Light Shading Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_background(hdr_cells[idx], "1F4E79")
        set_cell_margins(hdr_cells[idx], 120, 120, 150, 150)
        run = hdr_cells[idx].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        
    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate(row_data):
            row_cells[c_idx].text = text
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 80, 80, 150, 150)
            run = row_cells[c_idx].paragraphs[0].runs[0]
            run.font.size = Pt(9.5)
            if "Dirty" in text:
                run.font.bold = True
                run.font.color.rgb = ACCENT_WARN
            elif "Clean" in text:
                run.font.color.rgb = RGBColor(46, 125, 50) # Green
                
    # Section 3
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(24)
    h3.paragraph_format.space_after = Pt(6)
    run = h3.add_run("3. Data Quality & Anomalies Breakdown (The \"Dirty\" List)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "While profiling the files, I detected several anomalies that will cause issues down the road if we don't fix them. "
        "Here is exactly what I found:"
    )
    
    # Sub-heading 3.A
    h3a = doc.add_paragraph()
    h3a.paragraph_format.space_before = Pt(8)
    h3a.paragraph_format.space_after = Pt(4)
    run = h3a.add_run("A. Missing (Null) Values")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("fund_master.csv: ").bold = True
    bp1.add_run("There is ")
    bp1.add_run("1 missing value").bold = True
    bp1.add_run(" in the risk_grade column. It belongs to scheme 999999 (\"Test Missing NAV Fund\"). We should check if we can infer this or if it's just a dummy test record.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("investors.csv: ").bold = True
    bp2.add_run("Investor INV004 (\"Bharath V\") is ")
    bp2.add_run("missing an email address").bold = True
    bp2.add_run(". Since we might need this for communication or unique lookups, we should note it.")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("transactions.csv: ").bold = True
    bp3.add_run("Transaction TXN1004 is ")
    bp3.add_run("missing the units value").bold = True
    bp3.add_run(". Luckily, we have the total transaction amount and can impute this by dividing the amount by the NAV value on the transaction date once we join the tables.")
    
    # Sub-heading 3.B
    h3b = doc.add_paragraph()
    h3b.paragraph_format.space_before = Pt(12)
    h3b.paragraph_format.space_after = Pt(4)
    run = h3b.add_run("B. Duplicate Records")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("fund_master.csv: ").bold = True
    bp1.add_run("Found ")
    bp1.add_run("1 duplicate row").bold = True
    bp1.add_run(" for HDFC Top 100 Fund (amfi_code: 125497). It was written twice in the source.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("transactions.csv: ").bold = True
    bp2.add_run("Transaction TXN1001 was ")
    bp2.add_run("duplicated").bold = True
    bp2.add_run(" in the list. We need to drop this duplicate so we don't double-count sales or portfolio balances!")
    
    # Sub-heading 3.C
    h3c = doc.add_paragraph()
    h3c.paragraph_format.space_before = Pt(12)
    h3c.paragraph_format.space_after = Pt(4)
    run = h3c.add_run("C. Data Type and Format Inconsistencies")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Date Fields Read as Strings: ").bold = True
    bp1.add_run("Currently, the date columns in nav_history.csv, investors.csv, transactions.csv, and benchmark_history.csv are being read as raw text objects. I'll need to parse these into actual datetime objects on Day 2 to allow proper sorting, grouping, and chronological analysis.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("Mixed Date Formats: ").bold = True
    bp2.add_run("In nav_history.csv, a record for scheme 119551 was entered in the DD/MM/YYYY format, while everything else is in YYYY-MM-DD. This breaks basic string-based sorting and will cause date parsing to fail if we don't handle it with a flexible date parser.")
    
    # Section 4
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)
    run = h4.add_run("4. AMFI Code Integrity & Cross-Validation")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "The AMFI Scheme Code is the unique 5-to-6 digit number assigned by the Association of Mutual Funds in India. "
        "It is our natural key to join the static fund metadata (like risk ratings, fund houses, categories) with price history or transactions."
    )
    p = doc.add_paragraph(
        "I ran a relational check between our fund_master.csv unique codes and our nav_history.csv unique codes. I found two integrity issues:"
    )
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Orphan Master Code: ").bold = True
    bp1.add_run("AMFI Code 999999 (\"Test Missing NAV Fund\") exists in our fund master metadata but has ")
    bp1.add_run("no daily price entries").bold = True
    bp1.add_run(" in nav_history.csv.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("Orphan NAV History Code: ").bold = True
    bp2.add_run("AMFI Code 888888 has daily price values inside nav_history.csv but ")
    bp2.add_run("does not exist in fund_master.csv").bold = True
    bp2.add_run(".")
    
    # Warning Callout Box using Table
    callout = doc.add_table(rows=1, cols=1)
    callout.style = 'Table Grid'
    cell = callout.rows[0].cells[0]
    set_cell_background(cell, "FDF2E9") # Soft Amber/Peach background
    set_cell_margins(cell, 150, 150, 200, 200)
    
    # Rust Red border on left (simulate in text since direct borders can be verbose in docx XML)
    p_c = cell.paragraphs[0]
    p_c.paragraph_format.space_after = Pt(0)
    run_c_title = p_c.add_run("WARNING ON REFERENCE INTEGRITY:\n")
    run_c_title.font.bold = True
    run_c_title.font.color.rgb = ACCENT_WARN
    run_c_title.font.size = Pt(10.5)
    
    run_c_text = p_c.add_run(
        "If we run a basic INNER JOIN in our SQL staging tables later, these records will be silently dropped. "
        "If we run a LEFT JOIN on nav_history, the orphan codes will have null details. "
        "Tomorrow, I will write a cleanup routine to decide whether to prune them or populate dummy names for them."
    )
    run_c_text.font.size = Pt(10)
    
    # Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(8)
    
    # Section 5
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)
    run = h5.add_run("5. Live Ingestion Results from API (mfapi.in)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "I wrote a small API tool (live_nav_fetch.py) to connect to the open API at https://api.mfapi.in and grab the live and historical NAV record sets. "
        "The script was fully successful, fetched all 6 schemes, and saved them as clean, individual CSVs under data/raw/:"
    )
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("HDFC Top 100 Direct (125497)").bold = True
    bp1.add_run(" -> hdfc_top_100_direct_nav.csv (3,091 records)")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("SBI Bluechip (119551)").bold = True
    bp2.add_run(" -> sbi_bluechip_nav.csv (3,236 records)")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("ICICI Bluechip (120503)").bold = True
    bp3.add_run(" -> icici_bluechip_nav.csv (3,307 records)")
    
    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.add_run("Nippon Large Cap (118632)").bold = True
    bp4.add_run(" -> nippon_large_cap_nav.csv (3,298 records)")
    
    bp5 = doc.add_paragraph(style='List Bullet')
    bp5.add_run("Axis Bluechip (119092)").bold = True
    bp5.add_run(" -> axis_bluechip_nav.csv (3,565 records)")
    
    bp6 = doc.add_paragraph(style='List Bullet')
    bp6.add_run("Kotak Bluechip (120841)").bold = True
    bp6.add_run(" -> kotak_bluechip_nav.csv (3,301 records)")
    
    # Section 6
    h6 = doc.add_paragraph()
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(6)
    run = h6.add_run("6. What's Next (Tomorrow's Cleanup Plan)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "To get these files ready for our database and dashboards, my plan for Day 2 is to:"
    )
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.add_run("Deduplicate: ").bold = True
    bp1.add_run("Clean up the duplicate rows in fund_master.csv and transactions.csv.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.add_run("Harmonize Dates: ").bold = True
    bp2.add_run("Use a flexible parser to convert all date columns into standard YYYY-MM-DD datetime objects, fixing the mixed-format entry in NAV history.")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.add_run("Impute Missing Units: ").bold = True
    bp3.add_run("Look up the NAV for Transaction TXN1004 and calculate the missing unit count programmatically.")
    
    bp4 = doc.add_paragraph(style='List Bullet')
    bp4.add_run("Align Keys: ").bold = True
    bp4.add_run("Address the orphan codes (999999 and 888888) so that our database queries execute cleanly with full referential integrity.")
    
    p = doc.add_paragraph(
        "Overall, the setup and ingestion went really well today. The data has a few typical real-world flaws, but they are all very manageable. "
        "Let me know if you have any questions or want me to tweak the cleaning plan for tomorrow!"
    )
    
    # Save Document
    output_path = os.path.join("C:\\Users\\jibum\\OneDrive\\Desktop\\Bluestock Internship\\reports", "day1_data_quality_summary.docx")
    doc.save(output_path)
    print(f"Successfully generated styled Word Document at: {output_path}")

if __name__ == "__main__":
    create_report()
